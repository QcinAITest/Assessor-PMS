"""
Generates QCI PMS High-Level Requirements Document as a .docx file.
Run: python3 docs/generate_requirements_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ──────────────────────────────────────────────────────────
QCI_BLUE   = RGBColor(0x1A, 0x56, 0x9A)   # header / heading blue
QCI_LIGHT  = RGBColor(0xD5, 0xE8, 0xF0)   # table header bg
QCI_MID    = RGBColor(0xEF, 0xF6, 0xFB)   # alternating row bg
DARK_TEXT  = RGBColor(0x1F, 0x2D, 0x3D)
MID_TEXT   = RGBColor(0x44, 0x44, 0x44)
BORDER_CLR = "BBCFE0"
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)


# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color=BORDER_CLR):
    """Apply thin borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_col_width(table, col_idx, width_inches):
    """Force a column width (overrides auto-fit)."""
    for row in table.rows:
        cell = row.cells[col_idx]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(width_inches * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)


def add_page_number(paragraph):
    """Append page X of Y to a paragraph."""
    run = paragraph.add_run('Page ')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r = OxmlElement('w:r')
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    paragraph._p.append(r)

    run2 = paragraph.add_run(' of ')
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.text = ' NUMPAGES '
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    r2 = OxmlElement('w:r')
    r2.append(fldChar3)
    r2.append(instrText2)
    r2.append(fldChar4)
    paragraph._p.append(r2)


def heading(doc, text, level=1, color=QCI_BLUE):
    """Add a styled heading paragraph."""
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.bold = True
    run.font.size = Pt(14) if level == 1 else Pt(12)
    run.font.name = 'Arial'
    return p


def body(doc, text, bold_parts=None, space_after=6):
    """Add a normal body paragraph, optionally bolding specific substrings."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if bold_parts is None:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = 'Arial'
        run.font.color.rgb = MID_TEXT
    else:
        # Split on bold_parts and reconstruct
        remaining = text
        for bold_text in bold_parts:
            idx = remaining.find(bold_text)
            if idx == -1:
                continue
            before = remaining[:idx]
            after  = remaining[idx + len(bold_text):]
            if before:
                r = p.add_run(before)
                r.font.size = Pt(10.5); r.font.name = 'Arial'
                r.font.color.rgb = MID_TEXT
            rb = p.add_run(bold_text)
            rb.font.bold = True; rb.font.size = Pt(10.5)
            rb.font.name = 'Arial'; rb.font.color.rgb = MID_TEXT
            remaining = after
        if remaining:
            r = p.add_run(remaining)
            r.font.size = Pt(10.5); r.font.name = 'Arial'
            r.font.color.rgb = MID_TEXT
    return p


def bullet(doc, text, level=0):
    """Add a bullet list item."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Arial'
    run.font.color.rgb = MID_TEXT
    return p


def numbered(doc, text, num):
    """Add a numbered list item."""
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Arial'
    run.font.color.rgb = MID_TEXT
    return p


def make_table(doc, headers, rows, col_widths=None, alt_rows=True):
    """Create a styled table with a blue header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, 'D5E8F0')
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.color.rgb = QCI_BLUE

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = 'EFF6FB' if (alt_rows and r_idx % 2 == 0) else 'FFFFFF'
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.name = 'Arial'
            run.font.color.rgb = MID_TEXT

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            set_col_width(table, i, w)

    doc.add_paragraph()  # spacing after table
    return table


def divider(doc):
    """Add a subtle horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B8CCE4')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ── Build document ────────────────────────────────────────────────────────────

doc = Document()

# Page setup — A4
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Default styles ────────────────────────────────────────────────────────────
normal_style = doc.styles['Normal']
normal_style.font.name = 'Arial'
normal_style.font.size = Pt(10.5)

for h_level in (1, 2, 3):
    try:
        s = doc.styles[f'Heading {h_level}']
        s.font.name = 'Arial'
        s.font.color.rgb = QCI_BLUE
    except Exception:
        pass

# ── Footer: page numbers ──────────────────────────────────────────────────────
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_page_number(footer_para)
run_footer = footer_para.add_run('   |   QCI Assessor PMS — High-Level Requirements v1.1')
run_footer.font.size = Pt(8)
run_footer.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
run_footer.font.name = 'Arial'

# ═══════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════

# Blue accent bar at top (using paragraph border trick)
top_bar = doc.add_paragraph()
top_bar.paragraph_format.space_before = Pt(0)
top_bar.paragraph_format.space_after  = Pt(0)
pPr = top_bar._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
top_el = OxmlElement('w:top')
top_el.set(qn('w:val'), 'single')
top_el.set(qn('w:sz'), '36')   # thick bar
top_el.set(qn('w:space'), '1')
top_el.set(qn('w:color'), '1A569A')
pBdr.append(top_el)
pPr.append(pBdr)

# Spacer
for _ in range(6):
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(0)

# Organisation tag
org_p = doc.add_paragraph()
org_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
org_run = org_p.add_run('QUALITY COUNCIL OF INDIA')
org_run.font.name = 'Arial'
org_run.font.size = Pt(11)
org_run.font.bold = True
org_run.font.color.rgb = QCI_BLUE
org_run.font.all_caps = True

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
title_p.paragraph_format.space_before = Pt(8)
title_run = title_p.add_run('Assessor Performance\nManagement System')
title_run.font.name = 'Arial'
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = DARK_TEXT

# Subtitle
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
sub_p.paragraph_format.space_before = Pt(6)
sub_run = sub_p.add_run('High-Level Requirements Document')
sub_run.font.name = 'Arial'
sub_run.font.size = Pt(16)
sub_run.font.bold = False
sub_run.font.color.rgb = QCI_BLUE

# Spacer
for _ in range(4):
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(0)

# Meta table
meta_table = doc.add_table(rows=5, cols=2)
meta_table.style = 'Table Grid'
meta_pairs = [
    ('Version',      '1.1'),
    ('Date',         'April 2026'),
    ('Status',       'In Development'),
    ('Product',      'QCI Assessor PMS'),
    ('Prepared by',  'QCI Technology Team'),
]
for i, (label, value) in enumerate(meta_pairs):
    row = meta_table.rows[i]
    lc = row.cells[0]; vc = row.cells[1]
    set_cell_bg(lc, 'D5E8F0'); set_cell_bg(vc, 'FFFFFF')
    set_cell_borders(lc); set_cell_borders(vc)
    lp = lc.paragraphs[0]
    lr = lp.add_run(label)
    lr.bold = True; lr.font.size = Pt(10); lr.font.name = 'Arial'
    lr.font.color.rgb = QCI_BLUE
    vp = vc.paragraphs[0]
    vr = vp.add_run(value)
    vr.font.size = Pt(10); vr.font.name = 'Arial'; vr.font.color.rgb = MID_TEXT
    lp.paragraph_format.space_before = Pt(4); lp.paragraph_format.space_after = Pt(4)
    vp.paragraph_format.space_before = Pt(4); vp.paragraph_format.space_after = Pt(4)
set_col_width(meta_table, 0, 1.5)
set_col_width(meta_table, 1, 3.0)

# Page break after cover
doc.add_paragraph().add_run().add_break(
    __import__('docx.enum.text', fromlist=['WD_BREAK_TYPE']).WD_BREAK_TYPE.PAGE
)


# ═══════════════════════════════════════════════════════
# SECTION 1 — Background & Problem Statement
# ═══════════════════════════════════════════════════════

heading(doc, '1. Background & Problem Statement')
divider(doc)

body(doc,
    'QCI oversees four national accreditation boards — NABL, NABH, NABCB, and NABET — '
    'each of which deploys teams of external assessors to evaluate accreditation-seeking '
    'organisations. Currently, assessor performance feedback is collected manually (paper or '
    'email), scored inconsistently across boards, and provides no longitudinal view of an '
    'assessor\'s improvement over time.',
    space_after=10
)

heading(doc, 'Key Pain Points', level=2)
pain_points = [
    'No standardised digital form for assessor evaluation across boards',
    'No single dashboard to compare assessor performance across NABL, NABH, NABCB, and NABET',
    'Assessors and staff data exist in QCI\'s HR/portal systems but must be manually re-entered into any new tool',
    'Board AMC members (admins) have no self-service tool to configure their own evaluation criteria',
]
for pp in pain_points:
    bullet(doc, pp)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
# SECTION 2 — Objectives
# ═══════════════════════════════════════════════════════

heading(doc, '2. Objectives')
divider(doc)

objectives = [
    ('Single configurable platform',
     'Serve all four boards with board-specific forms, roles, scoring models, and terminology from one system'),
    ('Data-driven performance tracking',
     'Per-audit scores, star ratings (1–5), and rolling cumulative ratings per assessor'),
    ('Frictionless feedback collection',
     'Public distributable form links — no assessor account or login required'),
    ('Automated form triggers',
     'Rules-based form generation triggered by audit completion, frequency, or events'),
    ('Portal integration',
     'Push-based sync from QCI\'s existing HR/portal systems — no manual data entry'),
    ('Strict data isolation',
     'A board admin can only see and manage their own board\'s data'),
]
for i, (bold_part, rest) in enumerate(objectives, 1):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(5)
    rb = p.add_run(bold_part + ': ')
    rb.bold = True; rb.font.size = Pt(10.5); rb.font.name = 'Arial'; rb.font.color.rgb = MID_TEXT
    rr = p.add_run(rest)
    rr.font.size = Pt(10.5); rr.font.name = 'Arial'; rr.font.color.rgb = MID_TEXT

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
# SECTION 3 — Stakeholders & User Roles
# ═══════════════════════════════════════════════════════

heading(doc, '3. Stakeholders & User Roles')
divider(doc)

make_table(
    doc,
    headers=['Role', 'Who', 'Responsibilities in PMS'],
    rows=[
        ['System Admin',         'QCI HQ IT team',                        'Creates boards, manages all users, runs bulk sync, views global audit log'],
        ['Board Admin\n(AMC Member)', 'Per-board Programme Management staff', 'Configures forms, roles, assessors, frequency rules for their board only'],
        ['Assessor',             'External domain expert on audits',       'Receives public form link; fills and submits evaluation (no account needed)'],
        ['QCI Portal',           'External HR/assessment system',          'Pushes assessor and user data into PMS via sync API endpoints'],
    ],
    col_widths=[1.4, 1.8, 4.0],
)

# ═══════════════════════════════════════════════════════
# SECTION 4 — Scope
# ═══════════════════════════════════════════════════════

heading(doc, '4. Scope')
divider(doc)

heading(doc, '4.1  In Scope', level=2)

scope_sections = {
    'Board Configuration': [
        'CRUD for boards (NABL, NABH, NABCB, NABET) with per-board config: rating engine, star bands, cumulative window, terminology',
        'Role mapping — translate generic internal role IDs (ROLE_LEAD, ROLE_PEER) to board-specific display labels',
        'Service lines and accreditation programs per board',
    ],
    'Form Builder': [
        'Visual form builder: add / edit / delete evaluation areas (top-level parameters) with percentage weights',
        'Sub-questions with response type selection: Rating 1–5, Yes/No, Percentage, Free Text',
        'Essential Criteria: mandatory YES/NO gates — any "No" flags the form regardless of numeric score',
        'Form versioning on every update; preview form before distribution',
    ],
    'Form Distribution & Collection': [
        'Generate shareable public links (token-based, no login required) per form per assessment',
        'Assessors fill forms via browser; auto-scored on submit',
        'Submission state machine: CREATED → SENT → PENDING → SUBMITTED / FLAGGED',
    ],
    'Scoring Engine': [
        'Per-form weighted score (parameters × weights, auto-normalised to 100%)',
        'Per-audit final score: weighted aggregation across all submitted forms for one assessor',
        'Cross-board normalised base-100 score for comparison',
        'Star rating (1–5) mapped from score via board-specific bands',
        'Cumulative rating: rolling average of last N audits (N configurable per board)',
        'Essential flag propagation: any flagged submission flags the audit score',
    ],
    'Frequency Rules': [
        'Configurable rules per (board, role, form): EVERY_AUDIT, POST_N_AUDITS, QUARTERLY, ANNUALLY, ON_EVENT',
        'Auto-generate pending form submissions when an assessment is marked complete',
    ],
    'Integration': [
        'Generic event ingestion endpoint (POST /api/v1/ingest/{board_code}) for external portals',
        'Assessment status check endpoint for portal polling',
        'Portal adapter layer: configurable role/event/vocabulary translation maps per portal per board',
        'Webhook registration for outbound events (SCORE_CALCULATED, ESSENTIAL_FLAGGED, etc.)',
        'Audit log for all INBOUND, OUTBOUND, and SYSTEM events',
    ],
    'Sync': [
        'Bulk assessor sync: upsert by employee_id, optional deactivate-missing, role translation via portal adapter',
        'Bulk admin user sync: upsert by email, auto-generated temporary passwords for new accounts',
    ],
    'Dashboard': [
        'Board Comparison Matrix: all boards side-by-side showing forms, roles, assessors, programs',
        'Per-board drill-down: assessor list, score history, cumulative ratings, form distribution',
    ],
    'Access Control': [
        'JWT-based auth with 8-hour tokens',
        'SYSTEM_ADMIN: unrestricted access across all boards',
        'BOARD_ADMIN: scoped strictly to own board — cannot view or modify other boards\' data',
        'Public form fill: token-based, no authentication required',
    ],
}

for section_title, items in scope_sections.items():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(section_title)
    r.bold = True; r.font.size = Pt(11); r.font.name = 'Arial'; r.font.color.rgb = QCI_BLUE
    for item in items:
        bullet(doc, item)

doc.add_paragraph()
heading(doc, '4.2  Out of Scope (v1)', level=2)
out_of_scope = [
    'Mobile native application',
    'Outbound webhook delivery (endpoint registered; dispatch not yet implemented)',
    'Assessor self-service portal (login, view own scores)',
    'Multi-language / regional language support',
    'Direct database integration with QCI portal (push model only — portal calls PMS)',
]
for item in out_of_scope:
    bullet(doc, item)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════
# SECTION 5 — Functional Requirements
# ═══════════════════════════════════════════════════════

heading(doc, '5. Functional Requirements')
divider(doc)

fr_sections = {
    'FR-1: Multi-Board Configuration': [
        'FR-1.1  Each board shall have an independent configuration profile (rating engine, star bands, cumulative window, stakeholder weights)',
        'FR-1.2  Each board shall define its own role taxonomy mapped to internal system role IDs',
        'FR-1.3  Board admins shall only access their assigned board; system admins access all boards',
    ],
    'FR-2: Form Builder': [
        'FR-2.1  Board admins shall create evaluation forms with a name, code, and stakeholder weight',
        'FR-2.2  Forms shall support a two-level parameter hierarchy (evaluation areas → sub-questions)',
        'FR-2.3  Sub-questions shall support five response types: Rating 1–5, Yes/No, Percentage, Free Text, Dropdown',
        'FR-2.4  Top-level parameters shall carry a percentage weight; the system shall auto-normalise if weights do not sum to 100%',
        'FR-2.5  Each form shall support Essential Criteria (mandatory YES/NO); a single "No" flags the entire form',
        'FR-2.6  Forms shall be versioned; any edit increments the version number',
        'FR-2.7  Board admins shall be able to preview a form before distribution',
    ],
    'FR-3: Form Distribution': [
        'FR-3.1  Board admins shall generate a unique shareable URL per form per assessment',
        'FR-3.2  Form fill shall require no assessor login — the token in the URL is the credential',
        'FR-3.3  A submitted token shall be non-reusable (re-opening shows "already submitted")',
    ],
    'FR-4: Scoring': [
        'FR-4.1  Form score shall be calculated automatically on submission',
        'FR-4.2  Final audit score shall aggregate multiple forms using their stakeholder weights',
        'FR-4.3  All scores shall be normalised to a 0–100 base for cross-board comparison',
        'FR-4.4  Star ratings (1–5) shall be derived from the board\'s configured score bands',
        'FR-4.5  Cumulative rating shall be a rolling average of the last N audits (N configurable; default 10)',
        'FR-4.6  Essential flags shall propagate from submission → audit score → cumulative rating',
    ],
    'FR-5: Frequency Rules': [
        'FR-5.1  The system shall support five trigger types: EVERY_AUDIT, POST_N_AUDITS, QUARTERLY, ANNUALLY, ON_EVENT',
        'FR-5.2  On assessment-complete, the system shall evaluate active rules for each evaluee and generate pending submissions',
        'FR-5.3  Each assessor shall have a running audit count used for POST_N_AUDITS evaluation',
    ],
    'FR-6: Integration': [
        'FR-6.1  External portals shall push assessment-complete events to the PMS',
        'FR-6.2  The PMS shall expose a status check endpoint that portals can poll before advancing workflow',
        'FR-6.3  Role and event translation between portal terminology and PMS shall be configurable per portal per board',
        'FR-6.4  All integration events shall be logged with direction, status, raw payload, and translated payload',
    ],
    'FR-7: Sync': [
        'FR-7.1  Assessor data shall be syncable in bulk using employee_id as the idempotent key',
        'FR-7.2  Admin user accounts shall be syncable in bulk using email as the idempotent key',
        'FR-7.3  Sync shall be idempotent — repeated calls with the same data produce no side effects',
        'FR-7.4  New users created via sync shall receive an auto-generated temporary password (shown once)',
        'FR-7.5  Sync shall never overwrite an existing user\'s password',
    ],
    'FR-8: Audit & Traceability': [
        'FR-8.1  All configuration changes shall be logged as SYSTEM-direction audit log entries',
        'FR-8.2  All inbound portal events shall be logged as INBOUND entries',
        'FR-8.3  System admins shall view the global audit log filtered by board, direction, and status',
    ],
}

for fr_title, items in fr_sections.items():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(fr_title)
    r.bold = True; r.font.size = Pt(11); r.font.name = 'Arial'; r.font.color.rgb = QCI_BLUE
    for item in items:
        bullet(doc, item)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
# SECTION 6 — Non-Functional Requirements
# ═══════════════════════════════════════════════════════

heading(doc, '6. Non-Functional Requirements')
divider(doc)

make_table(
    doc,
    headers=['Category', 'Requirement'],
    rows=[
        ['Security',        'Passwords stored as bcrypt hashes; JWTs signed with HS256; board data isolation enforced at the API layer on every request'],
        ['Performance',     'Board Comparison Matrix loads via parallel fetches (Promise.allSettled) — one failing board must not block others from displaying'],
        ['Availability',    'Local development: SQLite; production target: PostgreSQL with persistent storage and connection pooling'],
        ['Maintainability', 'Board-specific config stored in JSON columns — adding a new accreditation board requires no schema migration'],
        ['Scalability',     'Sync endpoints designed for bulk upsert of hundreds of assessors per call'],
        ['Usability',       'All forms accessible to assessors without a login; no mobile app required in v1'],
        ['Dark mode',       'Full dark mode support via Tailwind CSS dark: classes across all portal pages'],
    ],
    col_widths=[1.6, 5.6],
)

# ═══════════════════════════════════════════════════════
# SECTION 7 — Boards & Seed Data
# ═══════════════════════════════════════════════════════

heading(doc, '7. Boards & Initial Data')
divider(doc)

make_table(
    doc,
    headers=['Board', 'Rating Engine', 'Forms', 'Roles'],
    rows=[
        ['NABL',  'Numeric (1–5)',       '5 forms\n(F1_CAB, F2_LEAD, F3_PEER,\nF4_OFFICER, F5_COMMITTEE)', 'Lead Assessor, Peer Assessor, Technical Expert, Observer, Officer, Committee Member'],
        ['NABH',  'Percentage (0–100%)', '4 forms\n(F_HCO_PA, F_PEER,\nF_SECRETARIAT, F_COMMITTEE)',       'Principal Assessor, Co-Assessor, HCO Rep, Secretariat, Committee Member'],
        ['NABCB', 'Numeric (1–5)',       '1 form\n(F_CAB)',                                                'Team Leader, Assessor, Technical Expert, Trainee, Officer, Committee Member'],
        ['NABET', 'Numeric (1–5)',       '1 form\n(F_CLIENT)',                                             'Lead Assessor, Assessor, Technical Expert, Officer, Committee Member'],
    ],
    col_widths=[0.9, 1.4, 1.8, 3.1],
)

body(doc,
    'All four boards share 5 default Essential Criteria applied to every form: '
    'Professional Ethics, Confidentiality, Impartiality, Integrity, and Professional Conduct.',
    bold_parts=['Professional Ethics, Confidentiality, Impartiality, Integrity, and Professional Conduct.']
)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════
# SECTION 8 — Key Constraints & Decisions
# ═══════════════════════════════════════════════════════

heading(doc, '8. Key Design Decisions')
divider(doc)

make_table(
    doc,
    headers=['Decision', 'Rationale'],
    rows=[
        ['Push model for integration',          'PMS does not need to know QCI portal\'s internal API shape; aligns with the existing ingest endpoint pattern'],
        ['Token-based public forms\n(no assessor login)', 'Reduces friction; assessors are external and should not need portal accounts to submit feedback'],
        ['Soft-delete for assessors and users', 'Preserves historical scoring data; records are deactivated, never hard-deleted'],
        ['JSON config per board',               'Avoids sparse columns; adding a new board requires no database schema migration'],
        ['SQLite for dev / PostgreSQL for prod','Zero-setup local development; production-grade persistence via environment variable switch'],
    ],
    col_widths=[2.4, 4.8],
)

# ═══════════════════════════════════════════════════════
# SECTION 9 — Future Considerations
# ═══════════════════════════════════════════════════════

heading(doc, '9. Future Considerations')
divider(doc)

future = [
    'Assessor self-service portal — view own scores, download performance certificates',
    'Outbound webhook delivery — infrastructure registered; dispatch logic not yet implemented',
    'Email notifications to assessors when forms are distributed',
    'Role-based visibility within a board (e.g., AMC vs Secretariat view)',
    'Analytics and reporting dashboard — trend charts, board-level aggregates, export to PDF',
    'Multi-tenancy support beyond the 4 current QCI boards',
]
for item in future:
    bullet(doc, item)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════
# SECTION 10 — Glossary
# ═══════════════════════════════════════════════════════

heading(doc, '10. Glossary')
divider(doc)

make_table(
    doc,
    headers=['Term', 'Definition'],
    rows=[
        ['AMC',              'Assessment Management Committee — the administrative body per accreditation board'],
        ['Assessor',         'External domain expert who conducts accreditation assessments on behalf of a board'],
        ['Evaluee',          'The assessor being evaluated in a given form or assessment cycle'],
        ['Evaluator',        'The person filling in the evaluation form about the evaluee'],
        ['Essential Criterion', 'A mandatory YES/NO gate on a form; any "No" answer triggers a review flag regardless of numeric score'],
        ['Cumulative Rating','Rolling average of an assessor\'s last N audit scores (N is board-configured)'],
        ['Star Rating',      '1–5 star summary derived from a numeric score using the board\'s configured score bands'],
        ['PortalAdapter',    'Translation layer that maps an external portal\'s role/event IDs to PMS internal identifiers'],
        ['Submission Token', 'UUID embedded in a public form link; acts as the sole credential for form access (no login)'],
        ['Frequency Rule',   'Configuration record that determines when feedback forms are auto-generated for a given role/form combination'],
    ],
    col_widths=[1.8, 5.4],
)

# Footer note
note_p = doc.add_paragraph()
note_p.paragraph_format.space_before = Pt(20)
note_r = note_p.add_run('Document generated from product conversation history and implemented codebase — April 2026.')
note_r.font.size = Pt(8.5)
note_r.font.italic = True
note_r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
note_r.font.name = 'Arial'

# ── Save ──────────────────────────────────────────────────────────────────────
OUTPUT = '/Users/anoushka/Documents/qci-pms/docs/QCI_PMS_Requirements_v1.1.docx'
doc.save(OUTPUT)
print(f'Saved: {OUTPUT}')
